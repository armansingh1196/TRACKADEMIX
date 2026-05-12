from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parents[1]
DELIVERABLES_DIR = BASE_DIR / "deliverables"
ARTIFACTS_DIR = BASE_DIR / "artifacts"
OUTPUT_PATH = DELIVERABLES_DIR / "Student_Performance_AI_2_Page_Report.docx"
MODEL_SUMMARY_PATH = ARTIFACTS_DIR / "model_summary.json"
METRICS_PATH = ARTIFACTS_DIR / "metrics.json"


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(11)


def add_compact_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(item)
        run.font.size = Pt(9)


def add_image_row(document: Document, images: list[Path], width: Inches) -> None:
    table = document.add_table(rows=1, cols=len(images))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, image in enumerate(images):
        cell = table.cell(0, idx)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image), width=width)


def main() -> None:
    DELIVERABLES_DIR.mkdir(parents=True, exist_ok=True)
    summary = load_json(MODEL_SUMMARY_PATH)
    metrics = load_json(METRICS_PATH)

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("Student Performance AI\nTwo-Page Project Summary")
    run.bold = True
    run.font.size = Pt(16)

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(4)
    sub_run = sub.add_run("BIT Management System Prototype")
    sub_run.italic = True
    sub_run.font.size = Pt(10)

    intro_table = document.add_table(rows=1, cols=2)
    intro_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    intro_table.autofit = True

    left = intro_table.cell(0, 0)
    right = intro_table.cell(0, 1)
    set_cell_shading(left, "EEF4FF")
    set_cell_shading(right, "F8FAFC")

    p = left.paragraphs[0]
    r = p.add_run("Problem Statement\n")
    r.bold = True
    r.font.size = Pt(10)
    body = p.add_run(summary["problem_statement"])
    body.font.size = Pt(9)

    p = right.paragraphs[0]
    r = p.add_run("Current Results\n")
    r.bold = True
    r.font.size = Pt(10)
    for line in [
        f"Accuracy: {metrics['accuracy']:.4f}",
        f"Weighted F1: {metrics['weighted_f1']:.4f}",
        f"Train/Test: {metrics['train_rows']} / {metrics['test_rows']}",
        "Target: Low / Medium / High",
    ]:
        rr = p.add_run(f"{line}\n")
        rr.font.size = Pt(9)

    add_heading(document, "1. Objective")
    add_compact_bullets(
        document,
        [
            "Predict student performance bands from academic and behavioral inputs.",
            "Identify students who are academically at risk.",
            "Generate recommendation-oriented intervention support.",
        ],
    )

    add_heading(document, "2. Why Random Forest")
    add_compact_bullets(
        document,
        [
            "Works well on structured student data such as attendance, GPA, quizzes, and assignments.",
            "Captures non-linear relationships between academic indicators.",
            "More stable than a single decision tree and easier to explain in project review.",
            "Provides feature importance, which supports explainability.",
        ],
    )

    add_heading(document, "3. Inputs And Target")
    add_compact_bullets(
        document,
        [
            "Numeric features: attendance, assignments, quizzes, internal exam, labs, study hours, LMS logins, GPA, project score.",
            "Categorical features: department, study mode, internet access, part-time job, extracurricular activity.",
            f"Low: {summary['target_definition']['Low']}",
            f"Medium: {summary['target_definition']['Medium']}",
            f"High: {summary['target_definition']['High']}",
        ],
    )

    add_heading(document, "4. Working Of The System")
    add_compact_bullets(
        document,
        [
            "Dataset is generated or loaded.",
            "Numeric and categorical columns are preprocessed separately.",
            "Missing values are handled with median/most-frequent imputation.",
            "Categorical values are converted using one-hot encoding.",
            "Random Forest predicts the performance band.",
            "A recommendation layer converts weak indicators into academic actions.",
        ],
    )

    add_heading(document, "5. Core Visuals")
    add_image_row(
        document,
        [
            ARTIFACTS_DIR / "target_distribution.png",
            ARTIFACTS_DIR / "feature_importance.png",
        ],
        width=Inches(3.0),
    )

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    add_heading(document, "6. Model Performance")
    matrix = metrics["confusion_matrix"]
    labels = metrics["class_labels"]
    table = document.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Actual \\ Predicted", "Low", "Medium", "High"]
    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        cell.text = header
        set_cell_shading(cell, "DBEAFE")
    for row_index, label in enumerate(labels, start=1):
        table.cell(row_index, 0).text = label
        set_cell_shading(table.cell(row_index, 0), "F1F5F9")
        for col_index, value in enumerate(matrix[row_index - 1], start=1):
            table.cell(row_index, col_index).text = str(value)

    add_compact_bullets(
        document,
        [
            "The model predicts Medium-class students most strongly.",
            "Some Low and High students still overlap with Medium, which is realistic for a prototype.",
            "The current performance is suitable for demonstration and further extension with real data.",
        ],
    )

    add_heading(document, "7. Recommendation Output")
    add_compact_bullets(
        document,
        [
            "Low attendance triggers attendance mentoring.",
            "Weak internal exam score triggers remedial support.",
            "Low assignment or quiz scores trigger targeted academic guidance.",
            "Low study hours or LMS usage trigger engagement and planning recommendations.",
        ],
    )

    add_heading(document, "8. Risk Visualization")
    add_image_row(document, [ARTIFACTS_DIR / "risk_distribution.png"], width=Inches(5.9))

    add_heading(document, "9. How To Run")
    add_compact_bullets(
        document,
        [
            "Use run_app.bat or run_app.ps1 for one-click launch.",
            "Manual flow: create .venv, install requirements, generate data, train model, run Streamlit.",
            "Dashboard features: metrics, charts, sample predictions, and live student prediction form.",
        ],
    )

    add_heading(document, "10. Conclusion")
    add_compact_bullets(
        document,
        [
            "Random Forest is an appropriate choice because it is reliable, explainable, and effective on tabular educational data.",
            "The prototype combines prediction, analytics, and recommendations in a single dashboard.",
            "The next step is integration with real BIT Management System data.",
        ],
    )

    document.save(OUTPUT_PATH)
    print(f"Created {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
